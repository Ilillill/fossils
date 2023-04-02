from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Q
import json
from django.db.models import Min, Max, Avg
from django.core.mail import send_mail
from django.core.paginator import Paginator
from .forms import ImageForm
from django.http import HttpResponse

from .models import DBFossil, DBSpecies
from .forms import FormSpecies, FormFossil, FormEmail


@login_required
def home(request):
    user_species = DBSpecies.objects.filter(species_owner=request.user)
    fossils_without_specie = DBFossil.objects.filter(fossil_owner=request.user, fossil_species__isnull=True)   

    paginator = Paginator(user_species, 10)
    page = request.GET.get('page')
    species = paginator.get_page(page)

    return render(request, 'app_fossils/home.html', {
        'species': species,
        'fossils_without_specie': fossils_without_specie,
        "nv": "home",
    })


@login_required
def species_add(request):
    forms_species_add = FormSpecies(request.user)
    if request.method == "POST":
        forms_species_add = FormSpecies(request.user, request.POST, request.FILES)
        if forms_species_add.is_valid():
            new_species = forms_species_add.save(commit=False)
            new_species.species_owner = request.user
            new_species.save()
            messages.success(request, f"{new_species.species_name} added successfuly")
            return redirect('homepage')
    return render(request, "app_fossils/species_add.html", {'form_species_add': forms_species_add, "nv": "species_add",})


@login_required
def fossil_add(request):
    forms_fossil_add = FormFossil(request.user)
    if request.method == "POST":
        forms_fossil_add = FormFossil(request.user, request.POST, request.FILES)
        if forms_fossil_add.is_valid():
            new_fossil = forms_fossil_add.save(commit=False)
            new_fossil.fossil_owner = request.user
            new_fossil.save()
            forms_fossil_add.save_m2m()
            return redirect('fossil-selected', pk=new_fossil.id)
    return render(request, "app_fossils/fossils_add.html", {'forms_fossil_add': forms_fossil_add, "nv": "fossil_add",})


@login_required
def fossil_update(request, pk):
    fossil_to_update = get_object_or_404(DBFossil, id=pk)
    if fossil_to_update.fossil_owner != request.user:
        messages.error(request, "This data belongs to a different user.")
        return redirect('account_logout')
    if request.method == "POST":
        form_fossil_update = FormFossil(request.user, request.POST, request.FILES, instance=fossil_to_update)
        if form_fossil_update.is_valid():
            fossil_to_update = form_fossil_update.save(commit=False)
            fossil_to_update.fossil_owner = request.user
            fossil_to_update.save()
            messages.success(request, f"{fossil_to_update.fossil_name} updated successfuly")
            return redirect('fossil-selected', pk=fossil_to_update.id)
    else:
        form_fossil_update = FormFossil(request.user, instance=fossil_to_update)

    return render(request, 'app_fossils/fossils_update.html', {
        'form_fossil_update': form_fossil_update,
        'fossil_to_update': fossil_to_update,
        'nv': 'fossil_update'
    })


@login_required
def fossil_delete(request, pk):
    fossil_to_delete = get_object_or_404(DBFossil, id=pk)
    if fossil_to_delete.fossil_owner != request.user:
        messages.error(request, "This data belongs to a different user.")
        return redirect('account_logout')
    if request.method == "POST":
        fossil_to_delete.delete()
        messages.success(request, f"{fossil_to_delete.fossil_name} deleted successfully")
        if fossil_to_delete.fossil_species:
            return redirect('species-selected', pk=fossil_to_delete.fossil_species.id)
        else:
            return redirect('homepage')
    return render(request, 'app_fossils/fossils_delete.html', {'fossil_to_delete': fossil_to_delete,})


@login_required
def species(request, pk):
    specie = get_object_or_404(DBSpecies, id=pk)
    if request.user != specie.species_owner:
        messages.error(request, "This data belongs to a different user.")
        return redirect('account_logout')
    
    if request.method == "POST" and "species_update" in request.POST:
        form_species_update = FormSpecies(request.user, request.POST, request.FILES, instance=specie)
        if form_species_update.is_valid():
            specie = form_species_update.save(commit=False)
            specie.species_owner = request.user
            specie.save()
            messages.success(request, f"{specie.species_name} updated successfully")
            return redirect('species-selected', pk=pk)
    else:
        form_species_update = FormSpecies(request.user, instance=specie)

    if request.method == "POST" and "species_delete" in request.POST:
        specie.delete()
        messages.success(request, f"{specie.species_name} deleted successfully")
        return redirect('homepage')


    if 'reset' in request.GET:
        request.session.pop('current_sort', None)
        return redirect('species-selected', pk=pk)

    current_sort = request.session.get("current_sort") or "fossil_name"

    if "ordering" in request.GET:
        ordering = request.GET["ordering"]
        if ordering == current_sort:
            ordering = "-" + ordering
        request.session["current_sort"] = ordering
    else:
        ordering = current_sort

    fossils_in_species = DBFossil.objects.filter(fossil_species=specie).order_by(ordering)

    paginator = Paginator(fossils_in_species, 5)
    page = request.GET.get('page')
    fossils = paginator.get_page(page)

    return render(request, 'app_fossils/species.html', {
        'specie': specie,
        'fossils': fossils,
        'form_species_update': form_species_update,
        'ordering': ordering,
        'nv': 'species',
    })


@login_required
def fossil(request, pk):
    fossil = get_object_or_404(DBFossil, id=pk)
    events = DBEvent.objects.filter(fossils=fossil).distinct()

    if fossil.fossil_owner != request.user:
        messages.error(request, "This data belongs to a different user.")
        return redirect('account_logout')
    
    if request.method == "POST" and "fossil_delete" in request.POST:
        fossil.delete()
        messages.success(request, f"{fossil.fossil_name} deleted successfully")
        if fossil.fossil_species:
            return redirect('species-selected', pk=fossil.fossil_species.id)
        else:
            return redirect('homepage')
    
    if request.method == "POST" and "fossil_update" in request.POST:
        form_fossil_update = FormFossil(request.user, request.POST, request.FILES, instance=fossil)
        if form_fossil_update.is_valid():
            fossil = form_fossil_update.save(commit=False)
            fossil.fossil_owner = request.user
            fossil.save()
            messages.success(request, f"{fossil.fossil_name} updated successfuly")
            return redirect('fossil-selected', pk=fossil.id)
    else:
        form_fossil_update = FormFossil(request.user, instance=fossil)
    
    if request.method == "POST" and "fossil_favourite" in request.POST:
        if fossil.fossil_is_favourite:
            fossil.fossil_is_favourite = False
            messages.success(request, f"{fossil.fossil_name} removed from favourites")
        else:
            fossil.fossil_is_favourite = True
            messages.success(request, f"{fossil.fossil_name} added to favourites")
        fossil.save()

    if request.method == "POST" and "send_email" in request.POST:
        form_send_email = FormEmail(request.POST)
        if form_send_email.is_valid():

            text_subject = f'{fossil.fossil_name}'
            subject = form_send_email.cleaned_data['subject']
            if subject:
                text_subject = f'{fossil.fossil_name} - {subject}'

            text_price = ''
            price = form_send_email.cleaned_data['include_price']
            if price and fossil.fossil_value:
                text_price = f'<p>{fossil.fossil_value}</p>'

            text_note = ''
            note = form_send_email.cleaned_data['email_note']
            if note:
                text_note = f'<p>{note}</p>'

            text_description = f'{fossil.fossil_name} '
            if fossil.fossil_description_short:
                text_description += fossil.fossil_description_short
            if fossil.fossil_description_detailed:
                text_description += fossil.fossil_description_detailed

            html = f'''
                    <html>
                        <body>
                            <div>
                            <h1>{text_subject}</h1>
                            {text_price}
                            {text_note}
                            {text_description}
                            </div>
                        </body>
                    </html>'''
            
            from_email = "app@sharedprojects.co.uk"
            email = form_send_email.cleaned_data['email']
            to_email = (email, )
            try:
                send_mail(subject=text_subject, message="", from_email=from_email, recipient_list=to_email, html_message=html)
                messages.success(request, f"Email to {email} sent successfully.")
            except Exception as e:
                messages.error(request, f"Error, email not sent. {e}")

            return redirect('fossil-selected', pk=fossil.id)
        else:
            messages.error(request, f"Error, email not sent.")
            return redirect('fossil-selected', pk=fossil.id)
    else:
        form_send_email = FormEmail()

    return render(request, 'app_fossils/fossil.html', {
        'fossil': fossil,
        'form_fossil_update': form_fossil_update,
        'form_send_email': form_send_email,
        'events': events,
        'nv': 'fossil',
        })

##########################################################
##########    P R O F I L E    S E C T I O N    ##########
##########################################################

from .models import Profile
from .forms import FormAccount, FormProfile

@login_required
def user_profile(request):
    if request.method == "POST" and "django_user_model" in request.POST:
        form_django_user_model = FormAccount(request.POST, instance=request.user)
        if form_django_user_model.is_valid():
            form_django_user_model.save()
            return redirect('profile')
    else:
        form_django_user_model = FormAccount(instance=request.user)

    # Use Django get_or_create method to handle form submission depending whether the profile already exists or not
    user_profile, created = Profile.objects.get_or_create(user=request.user)

    if request.method == "POST" and "user_custom_profile" in request.POST:
        form_custom_profile = FormProfile(request.POST, request.FILES, instance=user_profile)
        if form_custom_profile.is_valid():
            user_profile = form_custom_profile.save(commit=False)
            user_profile.user = request.user
            user_profile.save()
            return redirect('profile')
    else:
        form_custom_profile = FormProfile(instance=user_profile)

    option = 'Update' if not created else 'Create'

    return render(request, 'app_fossils/user_profile.html', {
        "form_django_user_model": form_django_user_model,
        "form_custom_profile": form_custom_profile,
        "option": option,
        "nv": "user_profile",
        })


########################################################
##########    S E A R C H    S E C T I O N    ##########
########################################################


@login_required
def search_result_get(request):
    species = DBSpecies.objects.filter(species_owner=request.user)
    fossils = DBFossil.objects.filter(fossil_owner=request.user)

    query = request.GET.get('search_get')
    species_result, fossils_result = None, None
    if species:
        species_result = species.filter(Q(species_name__icontains=query.strip().lower()) | Q(species_description__icontains=query.strip().lower()))
    if fossils:
        fossils_result = fossils.filter(Q(fossil_name__icontains=query.strip().lower()) | Q(fossil_description_short__icontains=query.strip().lower()))

    return render(request, 'app_fossils/search_result.html', {
        'species_result': species_result,
        'fossils_result': fossils_result,
        'search_type': 'GET',
    })


@login_required
def search_result_post(request):
    species = DBSpecies.objects.filter(species_owner=request.user)
    fossils = DBFossil.objects.filter(fossil_owner=request.user)

    query = request.POST.get('search_post')
    species_result, fossils_result = None, None
    if species:
        species_result = species.filter(Q(species_name__icontains=query.strip().lower()) | Q(species_description__icontains=query.strip().lower()))
    if fossils:
        fossils_result = fossils.filter(Q(fossil_name__icontains=query.strip().lower()) | Q(fossil_description_short__icontains=query.strip().lower()))

    return render(request, 'app_fossils/search_result.html', {
        'species_result': species_result,
        'fossils_result': fossils_result,
        'search_type': 'POST',
    })


########################################################
##########    C H A R T S    S E C T I O N    ##########
########################################################

@login_required
def chart_simple(request):
    species = DBSpecies.objects.filter(species_owner=request.user)
    labels = [specie.species_name for specie in species]
    labels.append('NO SPECIE')
    data = [DBFossil.objects.filter(fossil_species=specie).count() for specie in species]
    data.append(DBFossil.objects.filter(fossil_owner=request.user, fossil_species__isnull=True).count())

    chart_data = {
        'labels': labels,
        'datasets': [{
            'label': 'Number of fossils in species',
            'data': data,
            'backgroundColor': '#00ADB5',
            'borderColor': '#EEEEEE',
            'pointBackgroundColor': 'white',
            'pointBorderColor': '#000000',
        }
        ]
    }

    return render(request, 'app_fossils/chart_simple.html', {'chart_data': json.dumps(chart_data), "nv": "chart_simple",})


@login_required
def chart_multi(request):
    species = DBSpecies.objects.filter(species_owner=request.user)
    fossil_ages = {'youngest': [], 'oldest': [],'average': []}

    for specie in species:
        fossils = DBFossil.objects.filter(fossil_species=specie)

        youngest_fossil = fossils.filter(fossil_age_in_years=fossils.aggregate(Min('fossil_age_in_years'))['fossil_age_in_years__min'])

        if youngest_fossil:
            fossil_ages['youngest'].append(float(youngest_fossil.first().fossil_age_in_years))
        else:
            fossil_ages['youngest'].append(None)

        oldest_fossil = fossils.filter(fossil_age_in_years=fossils.aggregate(Max('fossil_age_in_years'))['fossil_age_in_years__max'])
        if oldest_fossil:
            fossil_ages['oldest'].append(float(oldest_fossil.first().fossil_age_in_years))
        else:
            fossil_ages['oldest'].append(None)

        average_age = fossils.aggregate(Avg('fossil_age_in_years'))['fossil_age_in_years__avg']
        if average_age:
            fossil_ages['average'].append(float(average_age))
        else:
            fossil_ages['average'].append(None)

    labels = [specie.species_name for specie in species]

    chart_data = {
        'labels': labels,
        'datasets': [{
            'label': 'Oldest age',
            'data': fossil_ages['oldest'],
            'borderColor': '#0081C9',
            'backgroundColor': '#0081C9',
            'yAxisID': 'y',
        },
        {
            'label': 'Average age',
            'data': fossil_ages['average'],
            'borderColor': '#5BC0F8',
            'backgroundColor': '#5BC0F8',
            'yAxisID': 'y',
        },
        {
            'label': 'Youngest age',
            'data': fossil_ages['youngest'],
            'borderColor': '#86E5FF',
            'backgroundColor': '#86E5FF',
            'yAxisID': 'y',
        }
        ]
    }

    return render(request, 'app_fossils/chart_multi.html', {'chart_data': json.dumps(chart_data), "nv": "chart_multi",})

################################################################
##########    G A T H E R I N G S    S E C T I O N    ##########
################################################################

from .models import DBFossilGathering
from .forms import FormGathering

@login_required
def gatherings(request):
    user_gatherings = DBFossilGathering.objects.filter(gathering_owner=request.user)
    if request.method == "POST":
        form_gathering_add = FormGathering(request.user, request.POST, request.FILES)
        if form_gathering_add.is_valid():
            new_gathering = form_gathering_add.save(commit=False)
            new_gathering.gathering_owner = request.user
            new_gathering.save()
            messages.success(request, f"{new_gathering.gathering_name} added successfuly")
            return redirect('gathering-selected', pk=new_gathering.id)
    else:
        form_gathering_add = FormGathering(request.user)

    return render(request, "app_fossils/gatherings.html", {
        "user_gatherings": user_gatherings,
        "form_gathering_add": form_gathering_add,
        "nv": "gatherings"
    })


@login_required
def gathering_selected(request, pk):
    gathering = get_object_or_404(DBFossilGathering, id=pk)
    if gathering.gathering_owner != request.user:
        messages.error(request, "This data belongs to a different user.")
        return redirect('account_logout')

    fossils = gathering.gathered.all()
    if request.method == "POST" and "gathering_delete" in request.POST:
        gathering.delete()
        messages.success(request, f"{gathering.gathering_name} deleted successfully")
        return redirect('gatherings')
    
    if request.method == "POST" and "gathering_update" in request.POST:
        form_gathering_update = FormGathering(request.user, request.POST, request.FILES, instance=gathering)
        if form_gathering_update.is_valid():
            form_gathering_update.save(commit=False)
            form_gathering_update.gathering_owner = request.user
            form_gathering_update.save()
            messages.success(request, f"{gathering.gathering_name} updated successfuly")
            return redirect('gathering-selected', pk=gathering.id)
    else:
        form_gathering_update = FormGathering(request.user, instance=gathering)

    return render(request, "app_fossils/gathering.html", {
        "gathering": gathering,
        "fossils": fossils,
        "form_gathering_update": form_gathering_update,
        "nv": "gathering"
    })


########################################################
##########    E V E N T S    S E C T I O N    ##########
########################################################

from .models import DBEvent, FossilEvent
from .forms import FormEvent, FormFossilEvent, FormFossilReturn

@login_required
def events(request):
    events = DBEvent.objects.filter(event_owner=request.user)
    if request.method == "POST":
        form = FormEvent(request.user, request.POST)
        if form.is_valid():
            event = form.save(commit=False)
            event.event_owner = request.user
            event.save()
            messages.success(request, 'New event added successfully.')
            return redirect('event-selected', pk=event.id)
    else:
        form = FormEvent(request.user)
    return render(request, "app_fossils/events.html", {
        "events": events,
        "form": form,
        "nv": "events"
    })


@login_required
def event_selected(request, pk):
    event = get_object_or_404(DBEvent, id=pk)
    fossils = event.fossil_events.all()

    if request.method == "POST" and "form_delete" in request.POST:
        event.delete()
        messages.success(request, 'Fossil return date updated successfully.')
        return redirect('events')
    
    if request.method == "POST" and "form_update" in request.POST:
        form_event_update = FormEvent(request.user, request.POST, instance=event)
        if form_event_update.is_valid():
            form_event_update.save(commit=False)
            form_event_update.event_owner = request.user
            form_event_update.save()
            messages.success(request, 'Event updated successfully.')
            return redirect('event-selected', pk=pk)
    else:
        form_event_update = FormEvent(request.user, instance=event)

    if request.method == "POST" and "fossil_add" in request.POST:
        form_fossil_add = FormFossilEvent(request.user, request.POST)
        if form_fossil_add.is_valid():
            fossil_event = form_fossil_add.save(commit=False)
            fossil_event.event = event
            fossil_event.save()
            return redirect('event-selected', pk=pk)
    else:
        form_fossil_add = FormFossilEvent(request.user)

    return render(request, "app_fossils/event.html", {
        "form_event_update": form_event_update,
        "form_fossil_add": form_fossil_add,
        "event": event,
        "fossils": fossils,
    })


@login_required
def fossil_event_update(request, pk):
    fossil_event = get_object_or_404(FossilEvent, pk=pk)
    if request.method == 'POST' and "form_delete" in request.POST:
        fossil_event.delete()
        messages.success(request, 'Fossil event deleted successfully.')
        return redirect('event-selected', pk=fossil_event.event.pk)
    if request.method == 'POST' and "form_update" in request.POST:
        form_update = FormFossilEvent(request.user, request.POST, instance=fossil_event)
        if form_update.is_valid():
            form_update.save()
            messages.success(request, 'Fossil event updated successfully.')
            return redirect('event-selected', pk=fossil_event.event.pk)
    else:
        form_update = FormFossilEvent(request.user, instance=fossil_event)

    return render(request, 'app_fossils/fossil_event_update.html', {'form_update': form_update, "fossil_event": fossil_event})


@login_required
def fossil_event_return(request, pk):
    fossil_event = get_object_or_404(FossilEvent, pk=pk)
    if request.method == 'POST':
        form = FormFossilReturn(request.POST, instance=fossil_event)
        if form.is_valid():
            form.save()
            messages.success(request, 'Fossil return date updated successfully.')
            return redirect('event-selected', pk=fossil_event.event.pk)
    else:
        form = FormFossilReturn(instance=fossil_event)

    return render(request, 'app_fossils/fossil_event_return.html', {'form': form, "fossil_event": fossil_event})


#####################################
##########    T O O L S    ##########
#####################################


@login_required
def tools(request):
    return render(request, 'app_fossils/tools.html', {"nv": "tools"})


####################################################
##########    S E L E C T    I M A G E    ##########
####################################################


@login_required
def select_image(request):
    if request.method == 'POST':
        form = ImageForm(request.POST)
        if form.is_valid():
            selected_image = form.cleaned_data['selected_image']
            messages.success(request, f'You selected {selected_image}')
    else:
        form = ImageForm()
    return render(request, 'app_fossils/select_image.html', {'form': form})




def my_text(request):
    return HttpResponse('Hello, world!')

def my_html(request):
    html = '<html><body><h1>Hello, world!</h1></body></html>'
    return HttpResponse(html)


###########################################
##########    T I M E L I N E    ##########
###########################################


@login_required
def timeline(request):
    dates = DBFossilGathering.objects.filter(gathering_owner=request.user).order_by('-gathering_date')
    last_item = dates.last()
    div_height = 0
    if dates:
        div_height = dates.last().days_ago
    div_large = div_height+100

    return render(request, "app_fossils/timeline.html", {
        "dates": dates,
        'div_height': div_height,
        "div_large": div_large,
        'last_item': last_item,
        "nv":"timeline",
    })

import csv

@login_required
def print_csv(request):
    fossils = DBFossil.objects.filter(fossil_owner=request.user)
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="mymodel.csv"'
    writer = csv.writer(response)
    writer.writerow(['Name', 'Age', 'Value'])
    for obj in fossils:
        writer.writerow([obj.fossil_name, obj.fossil_age_mln, obj.fossil_value])
    return response

@login_required
def print_pdf(request):
    fossils = DBFossil.objects.filter(fossil_owner=request.user)
    return render(request, "app_fossils/print_pdf.html", {"fossils": fossils})

from docx import Document
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Mm
from docx.enum.style import WD_STYLE_TYPE
from django.utils.html import strip_tags
from docx.shared import RGBColor
import io
        
@login_required
def print_word(request):
    def print_fossil(selected_fossil):
        print(selected_fossil)
        fossil_name_text = document.add_paragraph(f'{selected_fossil.fossil_name}', style='title_style')
        fossil_name_text.paragraph_format.space_before = Pt(10)
        fossil_name_text.paragraph_format.space_after = Pt(0)
        fossil_name_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table_age_value = document.add_table(rows=2, cols=2)
        header_cells = table_age_value.rows[0].cells
        header_cells[0].text = f'Age'
        header_cells[0].paragraphs[0].style='big_header_style'
        header_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[0].paragraphs[0].paragraph_format.space_before = Pt(0)
        header_cells[0].paragraphs[0].paragraph_format.space_after = Pt(10)
        header_cells[1].text = f'Value'
        header_cells[1].paragraphs[0].style='big_header_style'
        header_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[1].paragraphs[0].paragraph_format.space_before = Pt(0)
        header_cells[1].paragraphs[0].paragraph_format.space_after = Pt(10)

        address_cells = table_age_value.rows[1].cells
        add_age = '-'
        if selected_fossil.fossil_age_mln:
            add_age = selected_fossil.fossil_age_mln
        address_cells[0].text = add_age
        address_cells[0].paragraphs[0].style='normal_style'
        address_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_cells[0].paragraphs[0].paragraph_format.space_before = Pt(0)
        address_cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
        add_value = '-'

        if selected_fossil.fossil_value:
            add_value = str(selected_fossil.fossil_value)
        address_cells[1].text = add_value
        address_cells[1].paragraphs[0].style='normal_style'
        address_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_cells[1].paragraphs[0].paragraph_format.space_before = Pt(0)
        address_cells[1].paragraphs[0].paragraph_format.space_after = Pt(0)

        if selected_fossil.fossil_description_short:
            tx = re.sub('[\t]+', ' ', strip_tags(selected_fossil.fossil_description_short))
            tx.replace('\n ', '\n').strip()
            note_par = document.add_paragraph('', style='normal_style')
            note_par.add_run(f'Note: ').bold=True
            note_par.add_run(f'{tx}').bold=False
            note_par.paragraph_format.space_before = Pt(5)
            note_par.paragraph_format.space_after = Pt(0)

        separator_line = document.add_paragraph('________________________________________________________________________________________________________________', style='normal_style')
        separator_line.alignment = WD_ALIGN_PARAGRAPH.CENTER


    fossils = DBFossil.objects.filter(fossil_owner=request.user)
    document = Document()

    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(5)
    section.right_margin = Mm(5)
    section.top_margin = Mm(5)
    section.bottom_margin = Mm(5)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    styles = document.styles

    style_title = styles.add_style('title_style', WD_STYLE_TYPE.PARAGRAPH)
    font_style_title = style_title.font
    font.name = 'Times New Roman'
    font_style_title.size = Pt(18)
    font_style_title.color.rgb = RGBColor(56, 118, 29)
    font_style_title.bold = True

    style_big_header = styles.add_style('big_header_style', WD_STYLE_TYPE.PARAGRAPH)
    font_style_big_header = style_big_header.font
    font_style_big_header.name = 'Times New Roman'
    font_style_big_header.size = Pt(14)
    font_style_big_header.bold = True

    style_normal = styles.add_style('normal_style', WD_STYLE_TYPE.PARAGRAPH)
    font_style_normal = style_normal.font
    font_style_normal.name = 'Times New Roman'
    font_style_normal.size = Pt(10)
    font_style_normal.bold = False

    for fossil in fossils:
        print_fossil(fossil)

    docx_title = f'Fossils.docx'
    doc_file = io.BytesIO()
    document.save(doc_file)
    length = doc_file.tell()
    doc_file.seek(0)
    response = HttpResponse(doc_file.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response

